import os
import uuid
from gtts import gTTS
from docx import Document
import PyPDF2
import re

def text_to_speech(text, lang_code="en", out_dir="tmp_audio"):
    """
    Convert text to speech and return mp3 path.
    Cleans text to avoid gTTS reading markdown or special chars literally.
    Supports multiple languages including Bengali.
    """
    # Remove markdown formatting (**bold**, __underline__, `code`)
    clean_text = re.sub(r"(\*\*|__|`)", "", text)
    # Remove other unwanted characters like excessive asterisks, underscores, tildes, hashtags
    clean_text = re.sub(r"[*_~#]", " ", clean_text)
    # Replace multiple spaces with single space
    clean_text = re.sub(r"\s+", " ", clean_text).strip()

    if not os.path.exists(out_dir):
        os.makedirs(out_dir, exist_ok=True)

    filename = f"{uuid.uuid4().hex}.mp3"
    path = os.path.join(out_dir, filename)
    tts = gTTS(text=clean_text, lang=lang_code)
    tts.save(path)
    return path

def extract_text_from_txt(file):
    """
    Extract text from a txt file. Handles Streamlit UploadedFile or local path.
    """
    try:
        if hasattr(file, "read"):
            file.seek(0)
            return file.read().decode("utf-8")
        else:
            with open(file, "r", encoding="utf-8") as f:
                return f.read()
    except Exception as e:
        raise RuntimeError(f"Failed to extract text from txt file: {e}")

def extract_text_from_docx(file):
    """
    Extract text from a docx file. Handles Streamlit UploadedFile or local path.
    """
    try:
        if hasattr(file, "read"):
            temp_path = f"tmp_{uuid.uuid4().hex}.docx"
            with open(temp_path, "wb") as f:
                f.write(file.read())
            doc = Document(temp_path)
            os.remove(temp_path)
        else:
            doc = Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        raise RuntimeError(f"Failed to extract text from docx file: {e}")

def extract_text_from_pdf(file):
    """
    Extract text from a pdf file. Handles Streamlit UploadedFile or local path.
    """
    try:
        if hasattr(file, "read"):
            temp_path = f"tmp_{uuid.uuid4().hex}.pdf"
            with open(temp_path, "wb") as f:
                f.write(file.read())
            reader = PyPDF2.PdfReader(temp_path)
            os.remove(temp_path)
        else:
            reader = PyPDF2.PdfReader(file)

        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text
    except Exception as e:
        raise RuntimeError(f"Failed to extract text from pdf file: {e}")
